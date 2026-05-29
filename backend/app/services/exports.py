from __future__ import annotations

from io import BytesIO
from typing import Any

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas


def build_pdf_export(assessment: dict[str, Any], report: dict[str, Any]) -> bytes:
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    y = height - 54

    def line(text: str, size: int = 10, gap: int = 16) -> None:
        nonlocal y
        if y < 72:
            pdf.showPage()
            y = height - 54
        pdf.setFont("Helvetica", size)
        pdf.drawString(54, y, text[:110])
        y -= gap

    line("CrisisMap Intelligence Report", 16, 24)
    line(f"Assessment: {assessment['name']}")
    line(f"Location: {assessment['location_name']}")
    line(f"Status: {assessment['status']} | Runtime: {assessment.get('runtime_seconds') or '--'}s")
    tokenrouter = assessment.get("tokenrouter") or report.get("tokenrouter") or {}
    line(f"AI: {tokenrouter.get('provider', 'unknown')} / {tokenrouter.get('model', 'unknown')}", 10, 24)

    line("Donor Summary", 13, 20)
    for chunk in _wrap(report.get("donor_summary", ""), 96):
        line(chunk)
    line("Damage Overview", 13, 20)
    for chunk in _wrap(report.get("damage_overview", ""), 96):
        line(chunk)

    line("Top Reconstruction Priorities", 13, 20)
    for item in report.get("priority_buildings", [])[:5]:
        line(f"#{item['rank']} {item['name']} - {item['status']} - score {item['damage_score']}")

    pdf.save()
    return buffer.getvalue()


def build_docx_export(assessment: dict[str, Any], report: dict[str, Any]) -> bytes:
    document = Document()
    document.add_heading("CrisisMap Intelligence Report", level=1)
    document.add_paragraph(f"Assessment: {assessment['name']}")
    document.add_paragraph(f"Location: {assessment['location_name']}")
    document.add_paragraph(f"Status: {assessment['status']}")
    tokenrouter = assessment.get("tokenrouter") or report.get("tokenrouter") or {}
    document.add_paragraph(f"AI: {tokenrouter.get('provider', 'unknown')} / {tokenrouter.get('model', 'unknown')}")

    document.add_heading("Donor Summary", level=2)
    document.add_paragraph(report.get("donor_summary", ""))
    document.add_heading("Damage Overview", level=2)
    document.add_paragraph(report.get("damage_overview", ""))
    document.add_heading("Top Reconstruction Priorities", level=2)
    for item in report.get("priority_buildings", [])[:5]:
        document.add_paragraph(
            f"#{item['rank']} {item['name']} - {item['status']} - score {item['damage_score']}",
            style="List Bullet",
        )

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _wrap(text: str, width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current: list[str] = []
    for word in words:
        candidate = " ".join([*current, word])
        if len(candidate) > width and current:
            lines.append(" ".join(current))
            current = [word]
        else:
            current.append(word)
    if current:
        lines.append(" ".join(current))
    return lines or [""]
